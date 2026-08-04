import streamlit as st
import google.generativeai as genai
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import json
import re

# Cấu hình trang Streamlit
st.set_page_config(page_title="Hệ thống Soạn Giáo án Tự Động 5512", layout="wide", page_icon="📝")

# ==========================================
# CẤU HÌNH GIAO DIỆN & TÙY CHỈNH STYLES
# ==========================================
st.markdown("""
    <style>
    .block-container {
        padding-top: 1.5rem !important;
        padding-bottom: 2rem !important;
    }

    html, body, [class*="css"] {
        font-family: 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
    }
    
    .sidebar-title {
        color: #1e293b;
        font-size: 21px !important;
        font-weight: 700;
        margin-bottom: 12px;
        display: flex;
        align-items: center;
        gap: 8px;
    }
    
    .step-header {
        color: #DC2626 !important;
        font-size: 21px !important;
        font-weight: 700 !important;
        margin-top: 15px !important;
        margin-bottom: 10px !important;
    }

    .custom-label {
        font-size: 21px !important;
        font-weight: 700 !important;
        color: #1E293B !important;
        margin-bottom: 4px !important;
        display: block !important;
    }
    
    div[data-testid="stWidgetLabel"] p {
        font-size: 20px !important;
        font-weight: 700 !important;
        color: #0F172A !important;
    }

    div[data-baseweb="input"] input, 
    div[data-baseweb="textarea"] textarea,
    div[data-baseweb="select"] span,
    div[data-baseweb="select"] div,
    ul[role="listbox"] li {
        font-size: 22px !important;
        font-weight: 600 !important;
        color: #1E3A8A !important;
    }

    .stButton button p {
        font-size: 21px !important;
        font-weight: 700 !important;
    }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# THANH BÊN (SIDEBAR) ĐĂNG NHẬP & CẤU HÌNH
# ==========================================
with st.sidebar:
    st.markdown("""
        <div class="sidebar-title">
            🔑 ĐĂNG NHẬP & CẤU HÌNH
        </div>
    """, unsafe_allow_html=True)
    
    api_key = st.text_input(
        "Google Gemini API Key:", 
        type="password", 
        placeholder="Dán mã AI Key vào đây...",
        help="Nhập API Key để kích hoạt trợ lý AI"
    )
    
    model_name = st.selectbox(
        "Mô hình AI xử lý:",
        ["gemini-3.6-flash", "gemini-3.6-flash"],
        index=0,
        help="Khuyên dùng 2.5-flash để phản hồi nhanh và chính xác"
    )
    
    st.markdown("---")
    
    st.markdown("""
        <div class="sidebar-title">
            👤 THÔNG TIN GIÁO VIÊN
        </div>
    """, unsafe_allow_html=True)
    
    school_name = st.text_input("Trường THPT:", "THPT NGUYỄN VĂN TRỖI")
    dept_name = st.text_input("Tổ chuyên môn:", "TỔ TOÁN")
    teacher_name = st.text_input("Họ và tên GV:", "Dương Tấn Tiến")
    
    st.markdown("---")
    st.caption("🟢 **Trạng thái:** Hệ thống sẵn sàng")

# ==========================================
# TIÊU ĐỀ ỨNG DỤNG
# ==========================================
st.markdown("""
    <div style="text-align: center; margin-bottom: 20px;">
        <div style="font-size: 37px; font-weight: 800; color: #1E293B; line-height: 1.2;">
            📝 HỆ THỐNG SOẠN KHBD (có tích hợp NLS, AI, STEM,...)
        </div>
        <div style="font-size: 23px; font-weight: 600; color: #2563EB; margin-top: 8px;">
            Tác giả: DƯƠNG TẤN TIẾN - GIÁO VIÊN TRƯỜNG THPT NGUYỄN VĂN TRỖI
        </div>
    </div>
""", unsafe_allow_html=True)

# ==========================================
# BƯỚC 1: THÔNG TIN BÀI DẠY
# ==========================================
st.markdown('<div class="step-header">📚 BƯỚC 1: THÔNG TIN MÔN HỌC & BÀI DẠY</div>', unsafe_allow_html=True)

col_sub, col_grd, col_dur = st.columns([2, 1, 1])
with col_sub:
    subject = st.selectbox("Môn học/Hoạt động GD:", ["Toán học", "Ngữ văn", "Tiếng Anh", "Vật lý", "Hóa học", "Sinh học", "Lịch sử", "Địa lý", "Tin học", "GDKT&PL", "Công nghệ"])
with col_grd:
    grade = st.selectbox("Khối lớp:", ["Lớp 10", "Lớp 11", "Lớp 12"])
with col_dur:
    duration = st.number_input("Số tiết:", value=2, min_value=1)

col_c, col_l = st.columns(2)
with col_c:
    chapter_title = st.text_input("Chương / Chủ đề:", placeholder="Ví dụ: Chương I. Mệnh đề và Tập hợp")
with col_l:
    lesson_title = st.text_input("Tên bài dạy:", placeholder="Ví dụ: Bài 1. Mệnh đề")

# ==========================================
# BƯỚC 2: CHỌN NGUỒN DỮ LIỆU DẠY HỌC (TỐI ƯU RÕ RÀNG)
# ==========================================
st.markdown('<div class="step-header">📂 BƯỚC 2: CUNG CẤP DỮ LIỆU SÁCH GIÁO KHOA & SÁCH GIÁO VIÊN</div>', unsafe_allow_html=True)

data_source_mode = st.radio(
    "Lựa chọn phương thức cung cấp dữ liệu:",
    ["📌 Cách 1: Tải File / Dán trực tiếp nội dung SGK & SGV (Chính xác 100%)", 
     "🔍 Cách 2: Tự động tra cứu danh mục bài học từ thư viện NXB"],
    index=0
)

uploaded_sgk = None
uploaded_sgv = None
sgk_text = ""
sgv_text = ""
requirements = ""

if "Cách 1" in data_source_mode:
    st.info("💡 **Khuyên dùng:** Dán văn bản trực tiếp hoặc tải file sẽ giúp AI lấy **chính xác 100% từng ví dụ, câu hỏi SGK và mục tiêu SGV**.")
    
    col_sgk_box, col_sgv_box = st.columns(2)
    with col_sgk_box:
        st.markdown("**📘 1. NỘI DUNG SÁCH GIÁO KHOA (Lấy ND bài học & Ví dụ):**")
        uploaded_sgk = st.file_uploader("Tải File SGK (PDF/Ảnh):", type=["pdf", "png", "jpg", "jpeg"], key="sgk_file")
        sgk_text = st.text_area("Hoặc DÁN VĂN BẢN SGK (Khởi động, HĐ, Ví dụ, Luyện tập):", height=150, placeholder="Copy các Hoạt động 1, 2, Ví dụ 1, 2 trong file PDF SGK dán vào đây...")

    with col_sgv_box:
        st.markdown("**📕 2. NỘI DUNG SÁCH GIÁO VIÊN (Lấy Mục tiêu):**")
        uploaded_sgv = st.file_uploader("Tải File SGV (PDF/Ảnh):", type=["pdf", "png", "jpg", "jpeg"], key="sgv_file")
        sgv_text = st.text_area("Hoặc DÁN MỤC TIÊU TỪ SGV (Mục tiêu bài học):", height=150, placeholder="Copy phần Mục tiêu (Kiến thức, Năng lực, Phẩm chất) trong SGV dán vào đây...")

else:
    # Cách 2: Tra cứu tự động
    if st.button("🔍 Cập nhật danh sách Chương & Bài học từ taphuan.nxbgd.vn", use_container_width=True):
        if not api_key:
            st.error("⚠️ Vui lòng nhập Gemini API Key ở thanh menu bên trái trước!")
        else:
            try:
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(model_name)
                
                prompt_fetch = f"""
                Hãy đóng vai Cơ sở dữ liệu chính thức của NXB Giáo dục Việt Nam (taphuan.nxbgd.vn).
                Liệt kê ĐẦY ĐỦ các Bài học thuộc môn {subject} - {grade} (Bộ sách Kết nối tri thức).
                Trả về dạng JSON mảng: [{"chapter": "...", "lesson": "...", "duration": 2, "req": "..."}]
                """
                with st.spinner("✨ Đang đồng bộ danh mục bài học..."):
                    res = model.generate_content(prompt_fetch)
                    raw_text = res.text.strip()
                    json_match = re.search(r'\[.*\]', raw_text, re.DOTALL)
                    clean_json = json_match.group(0) if json_match else raw_text
                    st.session_state['fetched_lessons'] = json.loads(clean_json)
                    st.success("🎉 Đã tải xong danh mục bài học!")
            except Exception as e:
                st.error(f"Lỗi khi tải danh mục: {e}")

    if 'fetched_lessons' in st.session_state and st.session_state['fetched_lessons']:
        lessons_data = st.session_state['fetched_lessons']
        lesson_titles = [f"{item['chapter']} - {item['lesson']}" for item in lessons_data]
        selected_idx = st.selectbox("👉 Chọn Bài học chuẩn từ danh sách:", range(len(lesson_titles)), format_func=lambda x: lesson_titles[x])
        current_item = lessons_data[selected_idx]
        chapter_title = current_item['chapter']
        lesson_title = current_item['lesson']
        duration = int(current_item['duration'])
        requirements = st.text_area("📌 Yêu cầu cần đạt:", value=current_item['req'], height=120)

# ==========================================
# BƯỚC 3: TÍCH HỢP NĂNG LỰC ĐẶC THÙ
# ==========================================
st.markdown('<div class="step-header">🚀 BƯỚC 3: TÍCH HỢP NĂNG LỰC ĐẶC THÙ</div>', unsafe_allow_html=True)

integrations = st.multiselect(
    "Lựa chọn yếu tố tích hợp:",
    [
        "Năng lực Số / Ứng dụng CNTT (Padlet, Kahoot, Geogebra...)", 
        "Tích hợp AI trong dạy và học (Gemini, ChatGPT, Canva...)", 
        "Giáo dục STEM / STEAM", 
        "Phát triển Tư duy phản biện", 
        "Học tập qua Dự án (PBL)"
    ],
    default=["Năng lực Số / Ứng dụng CNTT (Padlet, Kahoot, Geogebra...)", "Tích hợp AI trong dạy và học (Gemini, ChatGPT, Canva...)"]
)

# ==========================================
# HÀM XUẤT FILE WORD TỐI ƯU (ĐỊNH DẠNG CHÍNH XÁC, KHÔNG IN ĐẬM TÙY TIỆN)
# ==========================================
def format_paragraph_with_markdown(paragraph, text):
    """Phân tích chuỗi text chứa Markdown **tô đậm** và add vào paragraph chuẩn xác"""
    # Tách đoạn văn thành các chuỗi thường và các chuỗi nằm trong **...**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            if content:
                run = paragraph.add_run(content)
                run.bold = True
        else:
            if part:
                paragraph.add_run(part)

def generate_doc(content_text):
    doc = docx.Document()

    # Cấu hình lề trang chuẩn Công văn 5512 (2cm - 2cm - 3cm - 2cm)
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # Khung Thông tin Đầu trang (2 cột không viền)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.2)
    table.columns[1].width = Inches(3.2)

    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.line_spacing = 1.15
    run_school = p_left.add_run(f"TRƯỜNG: {school_name.upper()}\n")
    run_school.bold = True
    run_dept = p_left.add_run(f"TỔ: {dept_name.upper()}")
    run_dept.bold = True

    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.line_spacing = 1.15
    run_teacher_label = p_right.add_run("Họ và tên giáo viên:\n")
    run_teacher_label.bold = True
    run_teacher_val = p_right.add_run(teacher_name)
    run_teacher_val.bold = True

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>' % nsdecls('w'))
            tcPr.append(tcBorders)

    # Tiêu đề bài dạy
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(18)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(f"TÊN BÀI DẠY: {lesson_title.upper()}")
    r_title.bold = True
    r_title.font.size = Pt(14)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run(f"Môn học/Hoạt động giáo dục: {subject}; Lớp: {grade}\nThời gian thực hiện: ({duration} tiết)")
    r_sub.italic = True

    # ĐỌC VÀ XỬ LÝ ĐỊNH DẠNG VĂN BẢN TỪNG DÒNG
    lines = content_text.split('\n')
    for line in lines:
        line_str = line.strip()
        if not line_str or line_str.startswith("---") or line_str.startswith("# "):
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)

        # Lấy bản rõ không chứa markdown để kiểm tra tiêu đề
        clean_text = line_str.replace("**", "").replace("*", "")

        # 1. TIÊU ĐỀ MỤC LỚN (I. MỤC TIÊU, II. THIẾT BỊ..., III. TIẾN TRÌNH...) -> In đậm, size 14
        if clean_text.startswith("I. ") or clean_text.startswith("II. ") or clean_text.startswith("III. ") or clean_text.startswith("IV. "):
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(14)

        # 2. TIẾT VÀ HOẠT ĐỘNG (HOẠT ĐỘNG 1, HOẠT ĐỘNG 2...) -> In đậm, size 13
        elif clean_text.startswith("TIẾT ") or clean_text.startswith("HOẠT ĐỘNG ") or clean_text.startswith("Hoạt động "):
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(13)

        # 3. CÁC MỤC CON THUỘC MỤC I, II (1. Kiến thức, 2. Năng lực...) -> In đậm
        elif clean_text.startswith("1. ") or clean_text.startswith("2. ") or clean_text.startswith("3. ") or clean_text.startswith("4. "):
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(clean_text)
            run.bold = True

        # 4. TIẾN TRÌNH HOẠT ĐỘNG (a) Mục tiêu, b) Nội dung, c) Sản phẩm, d) Tổ chức thực hiện) -> In đậm
        elif clean_text.startswith("a)") or clean_text.startswith("b)") or clean_text.startswith("c)") or clean_text.startswith("d)"):
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.left_indent = Inches(0.15)
            run = p.add_run(clean_text)
            run.bold = True

        # 5. CÁC BƯỚC THỰC HIỆN (Bước 1:, Bước 2:, Bước 3:, Bước 4:) -> In đậm
        elif clean_text.startswith("Bước 1:") or clean_text.startswith("Bước 2:") or clean_text.startswith("Bước 3:") or clean_text.startswith("Bước 4:") or clean_text.startswith("- Bước 1:") or clean_text.startswith("- Bước 2:") or clean_text.startswith("- Bước 3:") or clean_text.startswith("- Bước 4:"):
            p.paragraph_format.left_indent = Inches(0.3)
            # Chỉ in đậm từ "Bước X: ..."
            format_paragraph_with_markdown(p, line_str)

        # 6. ĐẦU DÒNG THƯỜNG (- hoặc +) -> Thụt lề, tô đậm ĐÚNG TỪNG TỪ CÓ **
        elif clean_text.startswith("- ") or clean_text.startswith("+ "):
            p.paragraph_format.left_indent = Inches(0.3)
            format_paragraph_with_markdown(p, line_str)

        # 7. VĂN BẢN THƯỜNG KHÁC -> Tô đậm ĐÚNG TỪNG TỪ CÓ ** (Không bao giờ in đậm cả đoạn)
        else:
            format_paragraph_with_markdown(p, line_str)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# XỬ LÝ TẠO KHBD
# ==========================================
st.markdown("<br>", unsafe_allow_html=True)
if st.button("🚀 BẮT ĐẦU TẠO KHBD WORD CHUẨN 5512", type="primary", use_container_width=True):
    if not api_key:
        st.error("⚠️ Vui lòng nhập Google Gemini API Key ở thanh menu bên trái!")
    elif not lesson_title:
        st.error("⚠️ Vui lòng nhập Tên Bài dạy ở Bước 1!")
    else:
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)

            integration_str = ", ".join(integrations) if integrations else "Không"

            # Xây dựng ngữ cảnh dữ liệu SGK / SGV
            sgk_context = ""
            if sgk_text:
                sgk_context += f"\n[DỮ LIỆU SÁCH GIÁO KHOA DO GV CUNG CẤP]:\n{sgk_text}\n"
            if sgv_text:
                sgk_context += f"\n[DỮ LIỆU SÁCH GIÁO VIÊN DO GV CUNG CẤP]:\n{sgv_text}\n"

            prompt = f"""
            Bạn là Giáo viên Cốt cán giỏi nhất môn {subject}.
            Nhiệm vụ: Soạn Kế hoạch bài dạy (Giáo án) chuẩn 100% CÔNG VĂN 5512/BGDĐT.
            - Môn: {subject} ({grade})
            - Chương/Chủ đề: {chapter_title}
            - Bài dạy: {lesson_title}
            - Thời lượng: {duration} tiết
            - Yếu tố tích hợp: {integration_str}
            {sgk_context}

            YÊU CẦU NGHIÊM NGẠT VỀ NỘI DUNG:
            1. VỀ MỤC TIÊU (Mục I): Bắt buộc lấy chính xác từ dữ liệu Sách giáo viên (SGV) nếu có cung cấp ở trên.
            2. VỀ TIẾN TRÌNH DẠY HỌC (Mục III): Bắt buộc trích dẫn CHÍNH XÁC 100% các câu hỏi Khởi động, Hoạt động khám phá, Ví dụ, Luyện tập, Bài tập từ Sách giáo khoa (SGK) được cung cấp. KHÔNG tự bịa ví dụ hay bài tập khác.

            QUY ĐỊNH ĐỊNH DẠNG VĂN BẢN (TRÁNH LỖI IN ĐẬM TÙY TIỆN):
            - CHỈ dùng ký tự ** ** cho các từ khóa quan trọng cần nhấn mạnh. Tuyệt đối KHÔNG bao bọc toàn bộ câu hay toàn bộ đoạn văn bằng ** **.
            - Đảm bảo đúng cấu trúc chuẩn 5512:
              I. MỤC TIÊU
              1. Kiến thức:
              2. Năng lực:
              3. Phẩm chất:
              II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU
              1. Giáo viên:
              2. Học sinh:
              III. TIẾN TRÌNH DẠY HỌC
              HOẠT ĐỘNG 1: MỞ ĐẦU / KHỞI ĐỘNG
              a) Mục tiêu:
              b) Nội dung:
              c) Sản phẩm:
              d) Tổ chức thực hiện:
              - Bước 1: Chuyển giao nhiệm vụ
              - Bước 2: Thực hiện nhiệm vụ
              - Bước 3: Báo cáo, thảo luận
              - Bước 4: Kết luận, nhận định

              HOẠT ĐỘNG 2: HÌNH THÀNH KIẾN THỨC MỚI
              ...
              HOẠT ĐỘNG 3: LUYỆN TẬP
              ...
              HOẠT ĐỘNG 4: VẬN DỤNG
              ...
            """

            content_payload = [prompt]
            if uploaded_sgk is not None:
                content_payload.append({"mime_type": uploaded_sgk.type, "data": uploaded_sgk.getvalue()})
            if uploaded_sgv is not None:
                content_payload.append({"mime_type": uploaded_sgv.type, "data": uploaded_sgv.getvalue()})

            with st.spinner("✨ AI đang tổng hợp chính xác từ SGK & SGV để tạo Giáo án Word..."):
                response = model.generate_content(content_payload)

                st.success("🎉 Tạo giáo án thành công!")
                doc_file = generate_doc(response.text)
                
                st.download_button(
                    label="📥 TẢI FILE WORD GIÁO ÁN (.DOCX)",
                    data=doc_file,
                    file_name=f"KHBD_5512_{lesson_title.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

                st.markdown("---")
                st.markdown(response.text)

        except Exception as e:
            st.error(f"Đã có lỗi xảy ra: {e}")
